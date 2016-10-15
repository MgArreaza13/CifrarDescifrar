#este es el codigo que me permitira 
#abrir, leer, y editar un archivo Docx
#luego de instalar la libreria python-docx
#no se porque razon no logro leer archivos 
#docx al parecer se instalo mal el modulo

from docx import *

d = Document()

d.add_heading('Hamlet')
d.add_heading('dramatis personae' , 2)
d.add_paragrah('hola mundo')
d.save('hamlet.docx')


document = Document('hamlet.docx')
for para in document.paragraphs:
	print(para.txt)